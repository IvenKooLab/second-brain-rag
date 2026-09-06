"""Document loading: recursively scan directories for .md/.txt (and .pdf when
the optional pypdf package is installed), peel off Obsidian-style YAML
frontmatter (tags/aliases subset), collect [[wikilinks]], and expand chat-log
exports (ChatGPT / Claude `conversations.json`) into per-conversation docs."""
from __future__ import annotations

import hashlib
import re
from pathlib import Path

from loci.chatlog import CHATLOG_NAMES, parse_chatlog

try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import pymupdf4llm
    HAS_PDF_TABLES = True
except ImportError:
    HAS_PDF_TABLES = False

try:
    import docx as _docx  # python-docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

SUFFIXES = {".md", ".txt"}
_WIKILINK = re.compile(r"(?<!!)\[\[([^\]|#]+)(?:[#|][^\]]*)?\]\]")


def extract_wikilinks(body: str) -> str:
    """Collect [[wikilink]] targets as a comma-joined string (deduplicated).
    Obsidian embeds (![[...]]) are assets, not links — excluded."""
    seen: list[str] = []
    for m in _WIKILINK.finditer(body):
        target = m.group(1).strip()
        if target and target.lower() not in (t.lower() for t in seen):
            seen.append(target)
    return ",".join(seen)


def file_hash(content: str) -> str:
    return hashlib.sha1(content.encode("utf-8")).hexdigest()


def parse_frontmatter(content: str) -> tuple[dict, str]:
    """Parse a minimal YAML subset from leading --- fences: scalars and lists.
    Returns ({meta}, body); malformed fences are left in the body untouched."""
    if not content.startswith("---"):
        return {}, content
    lines = content.splitlines()
    if lines[0].strip() != "---":
        return {}, content
    end = None
    for i in range(1, len(lines)):
        if lines[i].strip() == "---":
            end = i
            break
    if end is None:
        return {}, content
    meta: dict = {}
    current_list: str | None = None
    for line in lines[1:end]:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("- ") and current_list:
            meta[current_list].append(stripped[2:].strip().strip('"\''))
            continue
        if ":" not in stripped:
            continue
        key, _, value = stripped.partition(":")
        key, value = key.strip(), value.strip()
        current_list = None
        if not value:
            meta[key] = []
            current_list = key
        elif value.startswith("[") and value.endswith("]"):
            inner = value[1:-1].strip()
            meta[key] = [v.strip().strip("\"'") for v in inner.split(",")] if inner else []
        else:
            meta[key] = value.strip("\"'")
    return meta, "\n".join(lines[end + 1:])


def tags_of(meta: dict) -> str:
    """Normalize the frontmatter `tags` key (list or string) to a comma-joined string."""
    raw = meta.get("tags", [])
    if isinstance(raw, str):
        items = [t.strip() for t in raw.split(",")]
    else:
        items = [str(t).strip() for t in raw]
    return ",".join(t for t in items if t)


def _read_text(p: Path) -> str | None:
    """Read a document as text/markdown; returns None when it can't be handled."""
    suffix = p.suffix.lower()
    if suffix in SUFFIXES:
        try:
            return p.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            print(f"[warn] not UTF-8, skipping: {p.name}")
            return None
    if suffix == ".pdf":
        if HAS_PDF_TABLES:
            try:
                # markdown output: headings, lists, and tables as pipe rows
                return pymupdf4llm.to_markdown(str(p))
            except Exception as e:
                print(f"[warn] pdf table extraction failed, falling back: {p.name} ({e})")
        if HAS_PDF:
            try:
                reader = PdfReader(str(p))
                return "\n\n".join((page.extract_text() or "") for page in reader.pages)
            except Exception as e:
                print(f"[warn] pdf unreadable, skipping: {p.name} ({e})")
                return None
        return None  # no pdf extra installed; `doctor` mentions it
    if suffix == ".docx":
        if not HAS_DOCX:
            return None  # silently skip; `doctor` mentions the optional extra
        try:
            document = _docx.Document(str(p))
            parts = [para.text for para in document.paragraphs if para.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n\n".join(parts)
        except Exception as e:
            print(f"[warn] docx unreadable, skipping: {p.name} ({e})")
            return None
    return None


def scan_sources(sources: list[dict]) -> list[dict]:
    """Return [{path, content, hash, tags, links, mtime, chunk_size?, chunk_overlap?}].
    Each [[sources]] entry may override `chunk_size` / `chunk_overlap`; absent
    keys mean "use the global [chunk] defaults"."""
    docs, seen = [], set()
    for src in sources:
        root = Path(src["path"]).expanduser()
        if not root.exists():
            print(f"[warn] directory not found, skipping: {root}")
            continue
        src_chunk_size = src.get("chunk_size")
        src_chunk_overlap = src.get("chunk_overlap")
        for p in sorted(root.rglob("*")):
            if not p.is_file():
                continue
            if p.name.lower() in CHATLOG_NAMES:
                ap = str(p.resolve())
                if ap in seen:
                    continue
                seen.add(ap)
                conversations = parse_chatlog(p)
                if conversations:
                    mtime = p.stat().st_mtime
                    seen_titles: set[str] = set()
                    for conv in conversations:
                        title, n = conv["title"], 2
                        while title.lower() in seen_titles:  # same-title convs would overwrite each other
                            title = f"{conv['title']} ({n})"
                            n += 1
                        seen_titles.add(title.lower())
                        conv["title"] = title
                        docs.append({"path": f"{ap}::{conv['title']}",
                                     "content": conv["text"],
                                     "hash": file_hash(conv["text"]),
                                     "tags": "chatlog", "links": "",
                                     "mtime": mtime,
                                     "chunk_size": src_chunk_size,
                                     "chunk_overlap": src_chunk_overlap})
                else:
                    print(f"[warn] unrecognized chat export, skipping: {p.name}")
                continue
            suffix = p.suffix.lower()
            if suffix not in SUFFIXES and suffix not in (".pdf", ".docx"):
                continue
            ap = str(p.resolve())
            if ap in seen:
                continue
            seen.add(ap)

            text = _read_text(p)
            if text is None or not text.strip():
                continue
            meta, body = parse_frontmatter(text)
            docs.append({"path": ap, "content": body, "hash": file_hash(text),
                         "tags": tags_of(meta), "links": extract_wikilinks(body),
                         "mtime": p.stat().st_mtime,
                         "chunk_size": src_chunk_size,
                         "chunk_overlap": src_chunk_overlap})
    return docs
