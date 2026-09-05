"""Office/PDF loader tests. These run only when the optional extras are
installed; CI installs `.[pdf]` and `.[docx]` so they run there too."""

import pytest

from loci import loaders

pytestmark_table = pytest.mark.skipif(not loaders.HAS_PDF_TABLES,
                                      reason="pymupdf4llm extra not installed")
pytestmark_docx = pytest.mark.skipif(not loaders.HAS_DOCX,
                                     reason="python-docx extra not installed")


def make_table_pdf(path):
    """Draw a small 2x3 grid table with PyMuPDF and return the pdf path."""
    import pymupdf
    doc = pymupdf.open()
    page = doc.new_page()
    rows = [["Item", "Qty"], ["Screws", "42"]]
    x0, y0, col_w, row_h = 72, 72, 100, 24
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            rect = pymupdf.Rect(x0 + c * col_w, y0 + r * row_h,
                                x0 + (c + 1) * col_w, y0 + (r + 1) * row_h)
            page.draw_rect(rect)
            page.insert_text((rect.x0 + 6, rect.y1 - 8), value, fontsize=11)
    doc.save(str(path))
    doc.close()
    return path


def test_pdf_tables_become_readable_text(tmp_path):
    pdf = make_table_pdf(tmp_path / "table.pdf")
    text = loaders._read_text(pdf)
    assert text, "pdf extraction returned nothing"
    assert "Screws" in text and "42" in text
    assert "Qty" in text


def test_pdf_table_survives_chunking_with_breadcrumb(tmp_path):
    from loci.chunker import split_markdown
    pdf = make_table_pdf(tmp_path / "table.pdf")
    text = loaders._read_text(pdf)
    chunks = split_markdown(text, 800, 100)
    joined = "\n".join(c["text"] for c in chunks)
    assert "Screws" in joined


def test_scan_sources_picks_up_pdf(tmp_path):
    make_table_pdf(tmp_path / "table.pdf")
    docs = loaders.scan_sources([{"path": str(tmp_path)}])
    assert len(docs) == 1 and docs[0]["path"].endswith("table.pdf")


def test_docx_paragraphs_and_tables(tmp_path):
    from docx import Document
    document = Document()
    document.add_heading("Quarterly Report", level=1)
    document.add_paragraph("Revenue grew in every region.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Region"
    table.cell(0, 1).text = "Revenue"
    table.cell(1, 0).text = "EMEA"
    table.cell(1, 1).text = "1.2M"
    p = tmp_path / "report.docx"
    document.save(str(p))

    text = loaders._read_text(p)
    assert "Quarterly Report" in text
    assert "Revenue grew in every region." in text
    assert "EMEA | 1.2M" in text  # table row pipe-joined


def test_scan_sources_picks_up_docx(tmp_path):
    from docx import Document
    document = Document()
    document.add_paragraph("hello docx")
    document.save(str(tmp_path / "a.docx"))
    docs = loaders.scan_sources([{"path": str(tmp_path)}])
    assert len(docs) == 1 and docs[0]["path"].endswith("a.docx")
